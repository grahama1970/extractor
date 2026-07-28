#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.10"
# dependencies = [
#   "typer>=0.12",
#   "python-docx>=1.0",
# ]
# ///
"""DOCX: numbering-based headings should be promoted with derived levels."""

from __future__ import annotations

import json
from pathlib import Path
import typer
from docx import Document

from extractor.pipeline.structured_pipeline import STRUCTURED_PIPELINES, run_structured_pipeline
from extractor.core.providers.docx import DOCXProvider

app = typer.Typer(add_completion=False)


@app.command()
def main(tmp_dir: Path = typer.Option(Path("data/results/structured_parity_smoke/docx_synth"))):
    """Generate a DOCX file with structured content and save to specified directory."""
    tmp_dir.mkdir(parents=True, exist_ok=True)
    docx_path = tmp_dir / "numbering.docx"

    doc = Document()
    doc.add_paragraph("4.1.5.4. Heading by numbering")
    doc.add_paragraph("A following paragraph")
    doc.save(docx_path)

    meta = STRUCTURED_PIPELINES[DOCXProvider]
    artifacts = run_structured_pipeline(
        DOCXProvider,
        docx_path,
        tmp_dir,
        stage_prefix=meta.stage_prefix,
        skip_export10=True,
        skip_embeddings10=True,
        fast_embeddings10=True,
    )
    s07 = json.loads(Path(artifacts["stage07"]).read_text())
    sections = s07.get("reflowed_sections") or []
    titles = [str((s or {}).get("title") or "").strip() for s in sections]
    if not any(t.startswith("4.1.5.4.") for t in titles):
        typer.echo(f"Numbering heading not promoted. Titles: {titles}", err=True)
        raise typer.Exit(code=1)
    typer.echo("DOCX numbering-based heading promotion passed.")


if __name__ == "__main__":
    app()
