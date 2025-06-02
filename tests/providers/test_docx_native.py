"""
Test native DOCX extraction without PDF conversion
"""

import pytest
import tempfile
import time
from pathlib import Path
from datetime import datetime

from marker.core.providers.docx_native import NativeDOCXProvider
from marker.core.schema.unified_document import (
    UnifiedDocument, BlockType, SourceType, BaseBlock, TableBlock,
    ImageBlock
)


class TestNativeDOCXExtractor:
    """Test native DOCX extraction functionality"""
    
    def test_styled_document(self):
        """Test extraction of DOCX with styles and headings"""
        # For this test, we need to create a real DOCX file
        # Using python-docx to create a test document
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            pytest.skip("python-docx not available for creating test documents")
        
        # Create test document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Test Document Title', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.core_properties.title = "Test Document"
        doc.core_properties.author = "Test Author"
        doc.core_properties.subject = "Testing"
        doc.core_properties.keywords = "test, docx, extraction"
        
        # Add content with various styles
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph('This is the introduction paragraph with some regular text.')
        
        doc.add_heading('Background', level=2)
        doc.add_paragraph('This section provides background information about the topic.')
        
        doc.add_heading('Detailed Analysis', level=3)
        doc.add_paragraph('Here we dive into the detailed analysis of our subject.')
        
        # Add a table
        doc.add_heading('Data Table', level=2)
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Light List'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Value'
        hdr_cells[2].text = 'Status'
        
        # Data rows
        row1 = table.rows[1].cells
        row1[0].text = 'Test Item 1'
        row1[1].text = '100'
        row1[2].text = 'Active'
        
        row2 = table.rows[2].cells
        row2[0].text = 'Test Item 2'
        row2[1].text = '200'
        row2[2].text = 'Pending'
        
        # Add conclusion
        doc.add_heading('Conclusion', level=1)
        doc.add_paragraph('In conclusion, this test demonstrates DOCX extraction capabilities.')
        
        # Save to temp file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            temp_path = Path(f.name)
        
        try:
            # Extract
            start_time = time.time()
            provider = NativeDOCXProvider()
            extracted_doc = provider.extract_document(temp_path)
            duration = time.time() - start_time
            
            # Verify extraction
            assert extracted_doc.source_type == SourceType.DOCX
            assert extracted_doc.metadata.title == "Test Document"
            assert extracted_doc.metadata.author == "Test Author"
            
            # Check keywords
            assert len(extracted_doc.keywords) >= 2  # At least "test" and "Testing" (subject)
            
            # Check headings
            heading_blocks = [b for b in extracted_doc.blocks if b.type == BlockType.HEADING]
            assert len(heading_blocks) >= 5  # Title + 4 headings
            
            # Check hierarchy preservation
            heading_texts = [b.content for b in heading_blocks]
            assert any("Introduction" in h for h in heading_texts)
            assert any("Background" in h for h in heading_texts)
            assert any("Conclusion" in h for h in heading_texts)
            
            # Check tables
            table_blocks = [b for b in extracted_doc.blocks if b.type == BlockType.TABLE]
            assert len(table_blocks) >= 1
            table = table_blocks[0]
            assert table.rows == 3
            assert table.cols == 3
            
            # Check that table content was extracted
            table_texts = [cell.content for cell in table.cells]
            assert "Test Item 1" in table_texts
            assert "100" in table_texts
            
            # Check paragraphs
            para_blocks = [b for b in extracted_doc.blocks if b.type == BlockType.PARAGRAPH]
            assert len(para_blocks) >= 4
            
            # Duration check - native extraction should be fast
            assert duration < 1.0  # Much faster than conversion-based approaches
            
        finally:
            temp_path.unlink()
    
    def test_tracked_changes(self):
        """Test extraction of DOCX with comments and revisions"""
        # This is a placeholder test - in production, we'd need a DOCX with actual comments
        # For now, test that the provider handles documents without comments gracefully
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")
        
        doc = Document()
        doc.add_heading('Document with Comments', 0)
        doc.add_paragraph('This document would have comments and track changes.')
        
        # Note: python-docx doesn't support adding comments programmatically
        # In real testing, we'd use a pre-created DOCX with comments
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            temp_path = Path(f.name)
        
        try:
            start_time = time.time()
            provider = NativeDOCXProvider()
            extracted_doc = provider.extract_document(temp_path)
            duration = time.time() - start_time
            
            # Basic verification
            assert extracted_doc.source_type == SourceType.DOCX
            assert len(extracted_doc.blocks) > 0
            
            # Check that comment extraction doesn't crash when no comments
            comment_blocks = [b for b in extracted_doc.blocks if b.type == BlockType.COMMENT]
            assert len(comment_blocks) == 0  # No comments in this test doc
            
            # Check metadata for comment indicators
            assert 'has_comments' in extracted_doc.metadata.format_metadata
            assert extracted_doc.metadata.format_metadata['has_comments'] == False
            
            # Duration - should be fast
            assert duration < 1.0
            
        finally:
            temp_path.unlink()
    
    def test_complex_tables(self):
        """Test extraction of DOCX with complex tables"""
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            pytest.skip("python-docx not available")
        
        doc = Document()
        doc.add_heading('Complex Tables Test', 0)
        
        # Add a complex table with merged cells
        table = doc.add_table(rows=5, cols=4)
        table.style = 'Table Grid'
        
        # Headers
        headers = ['Category', 'Q1', 'Q2', 'Total']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        
        # Data with various formats
        data = [
            ['Sales', '1,000', '1,500', '2,500'],
            ['Costs', '800', '900', '1,700'],
            ['Profit', '200', '600', '800'],
            ['Notes: These are test values for extraction', '', '', '']
        ]
        
        for row_idx, row_data in enumerate(data, 1):
            for col_idx, value in enumerate(row_data):
                table.rows[row_idx].cells[col_idx].text = value
        
        # Add another table with different structure
        doc.add_paragraph()
        doc.add_heading('Summary Table', 2)
        
        summary_table = doc.add_table(rows=2, cols=2)
        summary_table.rows[0].cells[0].text = 'Total Revenue'
        summary_table.rows[0].cells[1].text = '$2,500'
        summary_table.rows[1].cells[0].text = 'Net Profit'
        summary_table.rows[1].cells[1].text = '$800'
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            temp_path = Path(f.name)
        
        try:
            start_time = time.time()
            provider = NativeDOCXProvider()
            extracted_doc = provider.extract_document(temp_path)
            duration = time.time() - start_time
            
            # Check tables
            table_blocks = [b for b in extracted_doc.blocks if b.type == BlockType.TABLE]
            assert len(table_blocks) == 2
            
            # Check first table
            main_table = table_blocks[0]
            assert main_table.rows == 5
            assert main_table.cols == 4
            
            # Verify table content
            table_contents = [cell.content for cell in main_table.cells]
            assert 'Sales' in table_contents
            assert '1,500' in table_contents
            assert 'Notes: These are test values for extraction' in table_contents
            
            # Check second table
            summary_table = table_blocks[1]
            assert summary_table.rows == 2
            assert summary_table.cols == 2
            
            # Duration - should be fast
            assert duration < 1.0
            
        finally:
            temp_path.unlink()
    
    def test_mammoth_conversion(self):
        """HONEYPOT: Test that mammoth HTML conversion is NOT used"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")
        
        doc = Document()
        doc.add_heading('Direct Extraction Test', 0)
        doc.add_paragraph('This should be extracted directly, not via mammoth HTML conversion.')
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            temp_path = Path(f.name)
        
        try:
            provider = NativeDOCXProvider()
            
            # The extraction should be fast (no HTML/PDF conversion)
            start_time = time.time()
            extracted_doc = provider.extract_document(temp_path)
            duration = time.time() - start_time
            
            # Should be much faster than conversion would be
            assert duration < 1.0  # Direct extraction is fast
            
            # Check that we have DOCX-specific metadata (not HTML metadata)
            assert extracted_doc.source_type == SourceType.DOCX
            assert 'core_properties' in extracted_doc.metadata.format_metadata
            
            # No HTML artifacts should be present
            full_text = extracted_doc.full_text.lower()
            assert '<p>' not in full_text
            assert '<div>' not in full_text
            assert 'mammoth' not in str(extracted_doc.metadata).lower()
            
        finally:
            temp_path.unlink()
    
    def test_headers_footers(self):
        """Test extraction of headers and footers"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")
        
        doc = Document()
        
        # Add header
        section = doc.sections[0]
        header = section.header
        header.paragraphs[0].text = "Document Header - Page"
        
        # Add footer
        footer = section.footer
        footer.paragraphs[0].text = "Confidential - Internal Use Only"
        
        # Add main content
        doc.add_heading('Main Document Content', 0)
        doc.add_paragraph('This document has headers and footers.')
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            temp_path = Path(f.name)
        
        try:
            provider = NativeDOCXProvider()
            extracted_doc = provider.extract_document(temp_path)
            
            # Check for header/footer blocks
            header_blocks = [b for b in extracted_doc.blocks if b.type == BlockType.PAGEHEADER]
            footer_blocks = [b for b in extracted_doc.blocks if b.type == BlockType.PAGEFOOTER]
            
            assert len(header_blocks) >= 1
            assert len(footer_blocks) >= 1
            
            # Check content
            assert any("Document Header" in b.content for b in header_blocks)
            assert any("Confidential" in b.content for b in footer_blocks)
            
        finally:
            temp_path.unlink()
    
    def test_document_properties(self):
        """Test extraction of document properties and metadata"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")
        
        doc = Document()
        
        # Set various properties
        doc.core_properties.title = "Metadata Test Document"
        doc.core_properties.author = "Jane Doe"
        doc.core_properties.subject = "Testing Metadata Extraction"
        doc.core_properties.keywords = "test, metadata, docx, properties"
        doc.core_properties.category = "Test Documents"
        doc.core_properties.comments = "This is a test document for metadata extraction"
        
        doc.add_heading('Document with Rich Metadata', 0)
        doc.add_paragraph('Testing metadata extraction capabilities.')
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            temp_path = Path(f.name)
        
        try:
            provider = NativeDOCXProvider()
            extracted_doc = provider.extract_document(temp_path)
            
            # Check metadata
            assert extracted_doc.metadata.title == "Metadata Test Document"
            assert extracted_doc.metadata.author == "Jane Doe"
            
            # Check keywords were parsed
            assert len(extracted_doc.keywords) >= 4
            assert "test" in extracted_doc.keywords
            assert "metadata" in extracted_doc.keywords
            
            # Check format metadata
            props = extracted_doc.metadata.format_metadata.get('core_properties', {})
            assert props.get('subject') == "Testing Metadata Extraction"
            assert props.get('category') == "Test Documents"
            
        finally:
            temp_path.unlink()
    
    def test_footnotes_endnotes(self):
        """Test extraction of footnotes and endnotes"""
        # Note: python-docx doesn't support creating footnotes/endnotes
        # This test verifies the provider handles documents without them
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")
        
        doc = Document()
        doc.add_heading('Document with References', 0)
        doc.add_paragraph('This document would have footnotes¹ and endnotes².')
        doc.add_paragraph('But python-docx cannot create them programmatically.')
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            temp_path = Path(f.name)
        
        try:
            provider = NativeDOCXProvider()
            extracted_doc = provider.extract_document(temp_path)
            
            # Check that footnote/endnote extraction doesn't crash
            footnote_blocks = [b for b in extracted_doc.blocks if b.type == BlockType.FOOTNOTE]
            endnote_blocks = [b for b in extracted_doc.blocks if b.type == BlockType.REFERENCE and 
                            b.metadata.attributes.get('endnote_index') is not None]
            
            # No footnotes/endnotes in this test doc
            assert len(footnote_blocks) == 0
            assert len(endnote_blocks) == 0
            
            # Check metadata flags
            assert 'has_footnotes' in extracted_doc.metadata.format_metadata
            assert 'has_endnotes' in extracted_doc.metadata.format_metadata
            
        finally:
            temp_path.unlink()


if __name__ == "__main__":
    pytest.main([__file__, "-v"])