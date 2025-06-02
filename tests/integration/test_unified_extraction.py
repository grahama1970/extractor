"""
Test unified extraction across different file formats
Ensures all formats produce identical JSON schema
"""

import pytest
from pathlib import Path
import tempfile

from marker.core.providers.pdf import PdfProvider  
from marker.core.providers.html_native import NativeHTMLProvider
from marker.core.providers.docx_native import NativeDOCXProvider
from marker.core.schema.unified_document import UnifiedDocument, SourceType, BlockType


class TestUnifiedExtraction:
    """Test that all extractors produce the same JSON structure"""
    
    def test_schema_consistency(self):
        """Test that PDF, HTML, and DOCX produce same schema"""
        # Create test files (we'll use simple content for now)
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")
            
        # Create DOCX
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph.')
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            docx_path = Path(f.name)
            
        # Create HTML  
        html_content = """
        <html>
        <body>
            <h1>Test Document</h1>
            <p>This is a test paragraph.</p>
        </body>
        </html>
        """
        
        with tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w') as f:
            f.write(html_content)
            html_path = Path(f.name)
            
        try:
            # Extract with both providers
            docx_provider = NativeDOCXProvider()
            html_provider = NativeHTMLProvider()
            
            docx_doc = docx_provider.extract_document(docx_path)
            html_doc = html_provider.extract_document(html_path)
            
            # Both should produce UnifiedDocument
            assert isinstance(docx_doc, UnifiedDocument)
            assert isinstance(html_doc, UnifiedDocument)
            
            # Check source types
            assert docx_doc.source_type == SourceType.DOCX
            assert html_doc.source_type == SourceType.HTML
            
            # Both should have the same schema fields
            docx_fields = set(docx_doc.model_fields.keys())
            html_fields = set(html_doc.model_fields.keys())
            assert docx_fields == html_fields
            
            # Both should have blocks
            assert len(docx_doc.blocks) > 0
            assert len(html_doc.blocks) > 0
            
            # Both should have heading and paragraph blocks
            docx_types = {block.type for block in docx_doc.blocks}
            html_types = {block.type for block in html_doc.blocks}
            
            assert BlockType.HEADING in docx_types
            assert BlockType.HEADING in html_types
            assert BlockType.PARAGRAPH in docx_types
            assert BlockType.PARAGRAPH in html_types
            
            # Test JSON serialization
            docx_json = docx_doc.model_dump()
            html_json = html_doc.model_dump()
            
            # Check that JSON has same top-level keys
            assert set(docx_json.keys()) == set(html_json.keys())
            
            # Only source_type and file_type should differ
            assert docx_json['source_type'] == 'docx'
            assert html_json['source_type'] == 'html'
            
            # Metadata should have file_type field in format_metadata
            assert docx_doc.metadata.format_metadata.get('file_type') == 'docx'
            assert html_doc.metadata.format_metadata.get('file_type') == 'html'
            
        finally:
            docx_path.unlink()
            html_path.unlink()
            
    def test_arangodb_compatibility(self):
        """Test that extracted documents are ArangoDB-compatible"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")
            
        # Create test DOCX
        doc = Document()
        doc.add_heading('ArangoDB Test', 0)
        doc.add_paragraph('Testing ArangoDB compatibility.')
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            docx_path = Path(f.name)
            
        try:
            provider = NativeDOCXProvider()
            extracted_doc = provider.extract_document(docx_path)
            
            # Convert to JSON for ArangoDB
            json_doc = extracted_doc.model_dump()
            
            # Check that all required fields are present
            assert 'id' in json_doc
            assert 'source_type' in json_doc
            assert 'blocks' in json_doc
            assert 'metadata' in json_doc
            
            # Check that metadata has proper structure
            metadata = json_doc['metadata']
            assert 'format_metadata' in metadata
            assert metadata['format_metadata']['file_type'] == 'docx'
            
            # Ensure no None values at top level (ArangoDB requirement)
            for key, value in json_doc.items():
                if key in ['hierarchy', 'keywords', 'source_path', 'full_text', 'arango_key']:  # These can be None/empty
                    continue
                assert value is not None, f"Field {key} should not be None"
                
        finally:
            docx_path.unlink()


if __name__ == "__main__":
    pytest.main([__file__, "-v"])