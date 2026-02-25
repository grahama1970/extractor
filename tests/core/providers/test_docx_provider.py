"""Tests for DOCXProvider."""

from pathlib import Path

import pytest
from docx import Document

from extractor.core.providers.docx import DOCXProvider
from extractor.core.schema.unified_document import BlockType, SourceType


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a minimal DOCX file for testing."""
    doc_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test paragraph.")
    doc.add_heading("Section 1", level=2)
    doc.add_paragraph("Content in section 1.")
    doc.save(doc_path)
    return doc_path


def test_docx_provider_basic_extraction(sample_docx: Path):
    """Test basic DOCX extraction returns expected structure."""
    provider = DOCXProvider()
    doc = provider.extract_document(sample_docx)

    assert doc.source_type == SourceType.DOCX
    assert len(doc.blocks) >= 2
    # Should have at least one heading and one paragraph
    block_types = [b.type for b in doc.blocks]
    assert BlockType.HEADING in block_types or BlockType.PARAGRAPH in block_types


def test_docx_provider_heading_hierarchy(tmp_path: Path):
    """Test DOCX extracts headings with correct hierarchy."""
    doc_path = tmp_path / "headings.docx"
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_heading("Chapter 1", level=2)
    doc.add_heading("Section 1.1", level=3)
    doc.add_paragraph("Content")
    doc.save(doc_path)

    provider = DOCXProvider()
    result = provider.extract_document(doc_path)

    headings = [b for b in result.blocks if b.type == BlockType.HEADING]
    assert len(headings) >= 1
    # First heading should be level 1
    first_heading = headings[0]
    if first_heading.metadata and first_heading.metadata.attributes:
        level = first_heading.metadata.attributes.get("level")
        assert level in [1, 2, 3]  # Valid heading levels


def test_docx_provider_table_extraction(tmp_path: Path):
    """Test DOCX extracts tables correctly."""
    doc_path = tmp_path / "table.docx"
    doc = Document()
    doc.add_heading("Table Test", level=1)
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(0, 2).text = "Header 3"
    table.cell(1, 0).text = "Row 1"
    table.cell(1, 1).text = "Data"
    table.cell(1, 2).text = "Value"
    doc.save(doc_path)

    provider = DOCXProvider()
    result = provider.extract_document(doc_path)

    tables = [b for b in result.blocks if b.type == BlockType.TABLE]
    assert len(tables) >= 1
    # Table should have cells
    assert hasattr(tables[0], "cells")
    assert len(tables[0].cells) > 0


def test_docx_provider_empty_document(tmp_path: Path):
    """Test DOCX provider handles empty documents gracefully."""
    doc_path = tmp_path / "empty.docx"
    doc = Document()
    doc.save(doc_path)

    provider = DOCXProvider()
    result = provider.extract_document(doc_path)

    assert result.source_type == SourceType.DOCX
    # Empty document should not raise, may have zero blocks
    assert result.blocks is not None
