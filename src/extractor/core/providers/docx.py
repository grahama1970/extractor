"""
Module: docx.py
Purpose: Native DOCX extraction without PDF conversion

This module implements direct DOCX extraction using docx2python,
preserving all document features including comments, track changes,
styles, and complex formatting without lossy conversions.

External Dependencies:
- docx2python: https://github.com/ShayHill/docx2python

Example Usage:
>>> from extractor.core.providers.docx import DOCXProvider
>>> provider = DOCXProvider()
>>> document = provider.extract_document("example.docx")
>>> print(document.source_type)  # SourceType.DOCX
"""

import re
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
import os
from datetime import datetime

from docx2python import docx2python
from docx2python.iterators import iter_at_depth, iter_tables
from docx import Document as PythonDocxDocument
from loguru import logger
from extractor.core.providers.utils import emit_list_blocks, normalize_heading_level, ensure_hierarchy
from extractor.core.providers.fetcher_bridge import ensure_local_source, attach_fetcher_metadata

from extractor.core.schema.unified_document import (
    UnifiedDocument,
    BlockType,
    SourceType,
    BaseBlock,
    TableBlock,
    ImageBlock,
    BlockMetadata,
    DocumentMetadata,
    HierarchyNode,
    TableCell,
)


class DOCXProvider:
    """Direct DOCX extraction without PDF conversion"""

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.config = dict(config or {})
        self.block_counter = 0
        self.style_hierarchy = {
            "Title": 0,
            "Heading 1": 1,
            "Heading 2": 2,
            "Heading 3": 3,
            "Heading 4": 4,
            "Heading 5": 5,
            "Heading 6": 6,
            "Subtitle": 1,
        }
        # Regex for detecting numbered section patterns (e.g., "1.", "1.1", "1.1.1")
        self._section_number_re = re.compile(r'^(\d+\.)+\s*\S')
        # Numbering depth map (optional) used during list extraction
        self._numbering_map: Dict[str, List[int]] = {}

    def _detect_pdf_import(self, doc) -> bool:
        """
        Detect if DOCX was likely imported/converted from PDF.

        Signs of PDF-imported DOCX:
        - No defined custom styles (only built-ins like Normal)
        - All paragraphs use same generic style
        - Core properties may have converter metadata
        - Tables may have unusual cell structures

        Returns True if likely PDF-imported.
        """
        try:
            # Check core properties for converter hints
            props = doc.core_properties
            if props.creator:
                creator_lower = props.creator.lower()
                pdf_converters = ['adobe', 'pdf', 'acrobat', 'nitro', 'smallpdf', 'ilovepdf', 'zamzar']
                if any(conv in creator_lower for conv in pdf_converters):
                    return True

            # Check if document uses only generic styles
            style_names = set()
            for para in doc.paragraphs[:20]:  # Sample first 20 paragraphs
                try:
                    if para.style:
                        style_names.add(para.style.name)
                except Exception:
                    pass

            # If only "Normal" style or very few styles, likely PDF-imported
            if style_names <= {"Normal", "Body Text", "Default Paragraph Font"}:
                if len(doc.paragraphs) > 10:  # And has significant content
                    return True

        except Exception:
            pass

        return False

    def _detect_heading_heuristic(self, para, text: str) -> Optional[int]:
        """
        Heuristic heading detection for DOCX files imported from PDFs.

        These files often have broken/generic styles (e.g., "Normal" for everything).
        Uses visual formatting cues to detect headings.

        Returns heading level (1-6) or None if not a heading.
        """
        if not text or len(text) > 200:  # Headings are typically short
            return None

        # Check for numbered section patterns: "1.", "1.1", "1.1.1", etc.
        if self._section_number_re.match(text):
            # Count dots to estimate level
            dots = text.split()[0].count('.')
            return min(dots, 6)

        # Check formatting via paragraph runs
        try:
            runs = para.runs
            if not runs:
                return None

            # Check if entire paragraph is bold
            all_bold = all(run.bold for run in runs if run.text.strip())

            # Check font size (larger = heading)
            font_sizes = []
            for run in runs:
                if run.font and run.font.size:
                    # Size is in EMUs (914400 EMUs = 1 inch, 72 points = 1 inch)
                    pt_size = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size / 12700
                    font_sizes.append(pt_size)

            avg_size = sum(font_sizes) / len(font_sizes) if font_sizes else 0

            # Heuristics:
            # - Bold + short text = likely heading
            # - Large font (>14pt) + short = likely heading
            # - ALL CAPS + short = likely heading
            is_short = len(text) < 80
            is_all_caps = text.isupper() and len(text) > 3

            if is_short:
                if avg_size >= 18:
                    return 1  # Very large = level 1
                elif avg_size >= 14 or (all_bold and avg_size >= 12):
                    return 2  # Large or bold+medium = level 2
                elif all_bold:
                    return 3  # Bold alone = level 3
                elif is_all_caps:
                    return 2  # ALL CAPS = level 2

        except Exception:
            # If we can't access runs/formatting, skip heuristics
            pass

        return None

    # ------------------------------------------------------------------
    def _extract_simple(self, filepath: Path) -> UnifiedDocument:
        """Simple docx path for parity/gold fixtures.

        - Uses python-docx (not docx2python)
        - One block per paragraph, one block per table
        - No styling/metadata; deterministic and lightweight
        """
        from docx.oxml.ns import qn

        doc = PythonDocxDocument(filepath)
        blocks: List[BaseBlock] = []
        doc_id = filepath.stem

        # Detect if this DOCX was imported from PDF
        is_pdf_imported = self._detect_pdf_import(doc)
        if is_pdf_imported:
            logger.warning(
                f"DOCX file '{filepath.name}' appears to be imported from PDF. "
                "Styles and tables may be unreliable. Consider using the original PDF for better extraction accuracy."
            )

        # Build a mapping from XML elements to actual document paragraphs/tables
        # This preserves order while allowing proper style access
        para_map = {p._element: p for p in doc.paragraphs}
        table_map = {t._element: t for t in doc.tables}

        def iter_block_items(parent):
            """Iterate body elements in document order, yielding (type, element)."""
            for child in parent.iterchildren():
                if child.tag == qn("w:p"):
                    yield ("para", child)
                elif child.tag == qn("w:tbl"):
                    yield ("table", child)
                elif child.tag == qn("w:sdt"):
                    # content controls can wrap paragraphs/tables
                    for sub in child.iterchildren():
                        yield from iter_block_items(sub)

        for item_type, element in iter_block_items(doc.element.body):
            if item_type == "table":
                table = table_map.get(element)
                if not table:
                    continue
                rows: List[List[str]] = []
                for row in table.rows:
                    rows.append([cell.text for cell in row.cells])
                cells = []
                for r_idx, r in enumerate(rows):
                    for c_idx, cell in enumerate(r):
                        cells.append(TableCell(row=r_idx, col=c_idx, content=cell or ""))
                blocks.append(
                    TableBlock(
                        id=str(len(blocks)),
                        type=BlockType.TABLE,
                        content="table",
                        metadata=BlockMetadata(page_number=1),
                        rows=len(rows),
                        cols=len(rows[0]) if rows else 0,
                        cells=cells,
                    )
                )
            else:
                # Paragraph - look up in our map to get properly linked object
                para = para_map.get(element)
                if not para:
                    continue

                txt = para.text.strip()
                if not txt:
                    continue

                # Get style name safely from the properly linked paragraph
                style_name = ""
                try:
                    if para.style:
                        style_name = para.style.name or ""
                except Exception:
                    # Style might not be available (corrupted DOCX or imported from PDF)
                    pass

                block_type = BlockType.PARAGRAPH
                level = None

                if style_name and style_name in self.style_hierarchy:
                    block_type = BlockType.HEADING
                    level = self.style_hierarchy[style_name]
                elif style_name and style_name.startswith("Heading"):
                    # Handle "Heading 7", "Heading 8", etc.
                    block_type = BlockType.HEADING
                    try:
                        level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                else:
                    # Fallback: heuristic heading detection for PDF-imported DOCX
                    # These files often have broken/generic styles
                    heuristic_level = self._detect_heading_heuristic(para, txt)
                    if heuristic_level:
                        block_type = BlockType.HEADING
                        level = heuristic_level

                metadata = BlockMetadata(page_number=1)
                if level is not None:
                    metadata.attributes = {"level": level}

                blocks.append(
                    BaseBlock(
                        id=str(len(blocks)),
                        type=block_type,
                        content=txt,
                        metadata=metadata,
                    )
                )

        unified = UnifiedDocument(
            id=doc_id,
            source_type=SourceType.DOCX,
            blocks=blocks,
            metadata=DocumentMetadata(
                title=filepath.stem,
                author=None,
                format_metadata={
                    "pdf_imported": is_pdf_imported,
                    "pdf_import_warning": (
                        "This DOCX was likely converted from PDF. Styles and tables may be unreliable."
                        if is_pdf_imported else None
                    ),
                },
            ),
        )

        return ensure_hierarchy(unified)

    def extract_document(self, filepath: Union[str, Path]) -> UnifiedDocument:
        """Extract DOCX content to unified document format"""
        resolved_path, fetch_download = ensure_local_source(filepath)
        filepath = Path(resolved_path)
        logger.info(f"Extracting DOCX document: {filepath}")

        # Default: simple, parity-friendly path (one paragraph/table per block).
        # Opt out with DOCX_SIMPLE_MODE=0 to use the rich docx2python path.
        if os.environ.get("DOCX_SIMPLE_MODE", "1").lower() in {"1", "true", "yes"}:
            return self._extract_simple(filepath)

        # Optional: build numbering depth map (v2 opt-in)
        if os.environ.get("DOCX_USE_NUMBERING_DEPTH", "").lower() in {"1", "true"}:
            try:
                self._numbering_map = self._build_numbering_map(filepath)
            except Exception:
                self._numbering_map = {}

        # Extract with docx2python
        with docx2python(str(filepath)) as docx_content:
            # Extract all components
            metadata = self._extract_metadata(docx_content)
            attach_fetcher_metadata(metadata, fetch_download)
            blocks = self._extract_blocks(docx_content, filepath)

            # Post-process: promote numbered headings in mangled docs
            try:
                promoted = 0
                for b in blocks:
                    if b.type == BlockType.PARAGRAPH and isinstance(b.content, str):
                        t = b.content.strip()
                        if re.match(r"^\d+(?:\.\d+)*\.[\s\u00A0]+", t):
                            b.type = BlockType.HEADING
                            try:
                                dots = t.count('.')
                                b.metadata.attributes["level"] = max(1, min(dots + 1, 6))
                            except Exception:
                                b.metadata.attributes["level"] = 2
                            promoted += 1
                if promoted:
                    logger.debug(f"Promoted {promoted} DOCX paragraphs to headings by numbering heuristic")
            except Exception:
                pass

            # Add comments as blocks
            comment_blocks = self._extract_comments(docx_content)
            blocks.extend(comment_blocks)

            # Build hierarchy from styles
            hierarchy = self._build_hierarchy(blocks)

            # Extract images
            image_blocks = self._extract_images(docx_content)
            blocks.extend(image_blocks)

            # Apply Claude table merge analysis if enabled
            if self.config.get("enable_table_merge_analysis", True):
                blocks = self._analyze_and_merge_tables(blocks)

            # Create unified document
            doc = UnifiedDocument(
                id=self._generate_doc_id(filepath),
                source_type=SourceType.DOCX,
                source_path=str(filepath),
                blocks=blocks,
                hierarchy=hierarchy,
                metadata=metadata,
                full_text=self._extract_full_text(blocks),
                keywords=self._extract_keywords(docx_content),
            )

            logger.info(f"Extracted {len(blocks)} blocks from DOCX")
            return ensure_hierarchy(doc, default_title=filepath.stem)

    def _generate_doc_id(self, filepath: Path) -> str:
        """Generate unique document ID"""
        return hashlib.md5(str(filepath).encode()).hexdigest()

    def _extract_metadata(self, docx_content) -> DocumentMetadata:
        """Extract document metadata"""
        props = docx_content.core_properties

        metadata = DocumentMetadata()

        # Map core properties with null checks
        if props and hasattr(props, "get"):
            metadata.title = props.get("title", "") if props.get("title") else ""
            metadata.author = (
                props.get("creator", "") or props.get("lastModifiedBy", "") if props else ""
            )

            # Parse dates
            created = props.get("created")
            if created:
                try:
                    metadata.created_date = datetime.fromisoformat(created.replace("Z", "+00:00"))
                except (ValueError, AttributeError) as e:
                    logger.debug(f"Failed to parse created date '{created}': {e}")

            modified = props.get("modified")
            if modified:
                try:
                    metadata.modified_date = datetime.fromisoformat(modified.replace("Z", "+00:00"))
                except (ValueError, AttributeError) as e:
                    logger.debug(f"Failed to parse modified date '{modified}': {e}")

            # Language from properties
            metadata.language = props.get("language", "")

            # Store all properties as format metadata
            metadata.format_metadata = {
                "file_type": "docx",  # Add file type for consistency
                "core_properties": props,
                "has_comments": bool(docx_content.comments),
                "has_footnotes": bool(docx_content.footnotes),
                "has_endnotes": bool(docx_content.endnotes),
            }

        return metadata

    def _extract_blocks(self, docx_content, filepath: Path) -> List[BaseBlock]:
        """Extract all content blocks from DOCX"""
        blocks = []

        # Check if we have paragraph-level data with styles
        if hasattr(docx_content, "document_pars"):
            # Use paragraph-level extraction for better style detection
            blocks.extend(self._extract_blocks_with_styles(docx_content))
        else:
            # Fallback to basic extraction
            blocks.extend(self._extract_blocks_basic(docx_content))

        # Extract tables separately using iter_tables
        table_blocks = self._extract_tables_properly(docx_content)
        blocks.extend(table_blocks)
        blocks.extend(self._extract_tables_python_docx(filepath))

        # Process headers and footers
        if docx_content.header:
            for section_idx, header in enumerate(docx_content.header):
                header_blocks = self._process_header_footer(
                    header, BlockType.PAGEHEADER, f"section-{section_idx}"
                )
                blocks.extend(header_blocks)

        if docx_content.footer:
            for section_idx, footer in enumerate(docx_content.footer):
                footer_blocks = self._process_header_footer(
                    footer, BlockType.PAGEFOOTER, f"section-{section_idx}"
                )
                blocks.extend(footer_blocks)

        # Process footnotes
        if docx_content.footnotes:
            footnote_blocks = self._process_footnotes(docx_content.footnotes)
            blocks.extend(footnote_blocks)

        # Process endnotes
        if docx_content.endnotes:
            endnote_blocks = self._process_endnotes(docx_content.endnotes)
            blocks.extend(endnote_blocks)

        return blocks

    def _extract_blocks_with_styles(self, docx_content) -> List[BaseBlock]:
        """Extract blocks using paragraph-level style information"""
        blocks = []

        # Debug: Let's see what docx2python returns for body
        logger.debug(f"Body content type: {type(docx_content.body)}")
        logger.debug(
            f"Body content length: {len(docx_content.body) if hasattr(docx_content.body, '__len__') else 'N/A'}"
        )
        if docx_content.body and len(docx_content.body) > 0:
            logger.debug(f"First body element type: {type(docx_content.body[0])}")

        # Iterate through document paragraphs at depth 4
        pending_items: List[tuple[str, int]] = []
        current_list_type: Optional[str] = None
        last_heading_id: Optional[str] = None

        def flush_list():
            nonlocal pending_items, current_list_type
            if pending_items:
                list_block, li_blocks = emit_list_blocks(
                    items=pending_items,
                    list_type=current_list_type or "ul",
                    parent_id=last_heading_id or "",
                    id_prefix="docx-list",
                    start_index=self.block_counter,
                )
                blocks.append(list_block)
                blocks.extend(li_blocks)
                pending_items = []
                current_list_type = None
        for par_idx, par in enumerate(iter_at_depth(docx_content.document_pars, 4)):
            # Check if this is a Par object (from docx2python v3)
            if hasattr(par, "style") and hasattr(par, "runs"):
                # Extract text from runs
                text_parts = []
                for run in par.runs:
                    if hasattr(run, "text") and run.text:
                        text_parts.append(run.text)

                text = "".join(text_parts).strip()
                if not text:
                    continue

                # Get style
                style = par.style if hasattr(par, "style") else ""

                # Determine block type based on style first
                block_type = BlockType.PARAGRAPH
                level = None

                style_name = style or ""
                if style_name:
                    if style_name == "Title":
                        block_type = BlockType.HEADING
                        level = 1
                    elif style_name == "Subtitle":
                        block_type = BlockType.HEADING
                        level = 2
                    elif style_name.startswith("Heading"):
                        block_type = BlockType.HEADING
                        try:
                            lvl = style_name.replace("Heading", "").strip()
                            level = normalize_heading_level(int(lvl) if lvl.isdigit() else 1)
                        except Exception:
                            level = 1
                    elif "List" in style_name or "List Paragraph" in style_name:
                        # Style indicates list; treat paragraph as list item
                        block_type = BlockType.LISTITEM

                # Fallback content heuristics for mangled DOCX (no heading styles)
                if block_type == BlockType.PARAGRAPH:
                    # e.g., "4.1.5.4. Heading text" (allow NBSP)
                    if re.match(r"^\d+(?:\.\d+)*\.[\s\u00A0]+", text):
                        block_type = BlockType.HEADING
                        level = normalize_heading_level(text.count(".") + 1)
                    elif len(text) < 80 and text.isupper():
                        block_type = BlockType.HEADING
                        level = 2
                    # Detect ad-hoc list markers (•, -, – or 1.)
                    elif re.match(r"^(?:[\u2022\-\–]\s+|\d+[\.)]\s+)", text):
                        block_type = BlockType.LISTITEM

                metadata = BlockMetadata(
                    attributes={"style": style_name, "paragraph_index": par_idx}, confidence=1.0
                )

                if level is not None:
                    metadata.attributes["level"] = level

                if block_type == BlockType.HEADING:
                    flush_list()
                    hb = BaseBlock(
                        id=self._generate_block_id(),
                        type=BlockType.HEADING,
                        content=text,
                        metadata=metadata,
                    )
                    blocks.append(hb)
                    last_heading_id = hb.id
                elif block_type == BlockType.LISTITEM:
                    # Simple depth estimate from leading spaces / bullets
                    m = re.match(r"^(?P<bullet>[\u2022\-\–]|\d+[\.)])\s+", text)
                    depth = 1
                    if m:
                        # count spaces before text (docx2python may not preserve indent reliably; keep 1)
                        depth = 1
                        text = text[m.end() :]
                    list_t = "ol" if (m and m.group('bullet') and m.group('bullet')[0].isdigit()) else "ul"
                    # v2 opt-in: derive depth from python-docx numbering map when available
                    if self._numbering_map:
                        ilvls = self._numbering_map.get(text) or self._numbering_map.get(text.strip())
                        if ilvls:
                            try:
                                depth = int(ilvls.pop(0)) + 1
                                if not ilvls:
                                    self._numbering_map.pop(text, None)
                            except Exception:
                                pass
                    if current_list_type and current_list_type != list_t:
                        flush_list()
                    current_list_type = list_t
                    pending_items.append((text, depth))
                else:
                    flush_list()
                    blocks.append(
                        BaseBlock(
                            id=self._generate_block_id(),
                            type=block_type,
                            content=text,
                            metadata=metadata,
                        )
                    )
            elif isinstance(par, str) and par.strip():
                # Plain text paragraph (shouldn't happen with document_pars)
                blocks.append(self._process_paragraph(par, 0, par_idx))

        # Don't extract tables here - they're already extracted as paragraphs
        # Tables in docx2python are part of the document structure, not separate

        flush_list()
        return blocks

    def _build_numbering_map(self, filepath: Path) -> Dict[str, List[int]]:
        """Best-effort map from paragraph text to Word numbering ilvl (depth-1).

        Only used when DOCX_USE_NUMBERING_DEPTH=1. We prefer this lightweight
        approximation over restructuring providers; ties (duplicate text)
        are handled in FIFO order via list pops.
        """
        mapping: Dict[str, List[int]] = {}
        try:
            doc = PythonDocxDocument(str(filepath))
        except Exception:
            return mapping
        for p in getattr(doc, "paragraphs", []):
            txt = (getattr(p, "text", "") or "").strip()
            if not txt:
                continue
            try:
                pPr = p._p.pPr if getattr(p, "_p", None) is not None else None  # type: ignore[attr-defined]
                numPr = pPr.numPr if pPr is not None else None
                ilvl = numPr.ilvl.val if (numPr is not None and getattr(numPr, "ilvl", None) is not None) else None
                if ilvl is not None:
                    mapping.setdefault(txt, []).append(int(ilvl))
            except Exception:
                continue
        return mapping

    def _extract_blocks_basic(self, docx_content) -> List[BaseBlock]:
        """Basic extraction when style information is not available"""
        blocks = []
        body_content = docx_content.body

        for section_idx, section in enumerate(body_content):
            section_blocks = self._process_section(section, section_idx)
            blocks.extend(section_blocks)

        return blocks

    def _extract_tables_from_body(self, body_content) -> List[BaseBlock]:
        """Extract tables from document body"""
        blocks = []

        for section_idx, section in enumerate(body_content):
            if isinstance(section, list):
                for element_idx, element in enumerate(section):
                    # Check if this is really a table (list of lists with consistent structure)
                    if isinstance(element, list) and len(element) > 0:
                        # All rows should be lists for this to be a table
                        if all(isinstance(row, list) for row in element):
                            # Check that this isn't just a flat list of paragraphs
                            # Tables have multiple rows with multiple cells
                            if len(element) > 1 or (len(element) == 1 and len(element[0]) > 1):
                                table_block = self._process_table(element, section_idx, element_idx)
                                if table_block:
                                    blocks.append(table_block)

        return blocks

    def _process_section(self, section: List, section_idx: int) -> List[BaseBlock]:
        """Process a document section"""
        blocks = []

        for element_idx, element in enumerate(section):
            if isinstance(element, list):
                # This is a table
                table_block = self._process_table(element, section_idx, element_idx)
                if table_block:
                    blocks.append(table_block)
            elif isinstance(element, str):
                # This is a paragraph
                if element.strip():
                    para_block = self._process_paragraph(element, section_idx, element_idx)
                    if para_block:
                        blocks.append(para_block)
            elif hasattr(element, "__iter__"):
                # Handle nested structures
                for item in element:
                    if isinstance(item, str) and item.strip():
                        para_block = self._process_paragraph(item, section_idx, element_idx)
                        if para_block:
                            blocks.append(para_block)

        return blocks

    def _process_paragraph(self, text: str, section_idx: int, para_idx: int) -> Optional[BaseBlock]:
        """Process a paragraph"""
        text = text.strip()
        if not text:
            return None

        # Try to detect heading style from content patterns
        block_type = BlockType.PARAGRAPH
        level = None

        # Common heading patterns
        if re.match(r"^(Chapter|Section|Part)\s+\d+", text, re.IGNORECASE):
            block_type = BlockType.HEADING
            level = 1
        elif re.match(r"^\d+\.\s+[A-Z]", text):  # Numbered heading
            block_type = BlockType.HEADING
            level = text.count(".") + 1
        elif len(text) < 100 and text.isupper():  # Short all-caps likely heading
            block_type = BlockType.HEADING
            level = 2

        block_id = self._generate_block_id()

        metadata = BlockMetadata(
            attributes={"section_index": section_idx, "paragraph_index": para_idx}, confidence=0.95
        )

        if level:
            metadata.attributes["level"] = level

        return BaseBlock(id=block_id, type=block_type, content=text, metadata=metadata)

    def _extract_tables_properly(self, docx_content) -> List[BaseBlock]:
        """Extract tables using docx2python's table iterator"""
        blocks = []

        # docx2python puts all content in a table-like structure
        # Real tables can be found by checking the lineage
        table_idx = 0

        # Get all tables from the body
        for table in iter_tables(docx_content.body):
            # Check if this is actually a table (not just the document structure)
            # Real tables will have multiple rows and/or multiple columns
            if len(table) > 1 or (len(table) == 1 and len(table[0]) > 1):
                # This looks like a real table
                table_block = self._process_table(table, 0, table_idx)
                if table_block:
                    blocks.append(table_block)
                    table_idx += 1

        return blocks

    def _extract_tables_python_docx(self, filepath: Path) -> List[BaseBlock]:
        """Fallback table extraction using python-docx when docx2python misses tables."""
        try:
            document = PythonDocxDocument(str(filepath))
        except Exception:
            return []

        blocks: List[BaseBlock] = []
        for table_index, table in enumerate(document.tables):
            rows = len(table.rows)
            cols = len(table.columns) if table.columns else 0
            cells: List[TableCell] = []
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    cells.append(TableCell(row=r_idx, col=c_idx, content=text))
            block_id = self._generate_block_id()
            blocks.append(
                TableBlock(
                    id=block_id,
                    type=BlockType.TABLE,
                    content={"title": None},
                    rows=rows,
                    cols=cols,
                    cells=cells,
                    headers=[0] if rows else None,
                    metadata=BlockMetadata(
                        attributes={"table_index": table_index, "source": "python-docx"},
                        confidence=0.9,
                    ),
                )
            )
        return blocks

    def _process_table(
        self, table_data: List, section_idx: int, table_idx: int
    ) -> Optional[TableBlock]:
        """Process a table"""
        if not table_data or not any(table_data):
            return None

        cells = []
        headers = []

        for row_idx, row in enumerate(table_data):
            if not isinstance(row, list):
                continue

            # Check if this is a header row (first row often is)
            if row_idx == 0 and all(isinstance(cell, str) for cell in row):
                # Simple heuristic: if all cells are short text, likely headers
                if all(len(str(cell).strip()) < 50 for cell in row if cell):
                    headers.append(row_idx)

            for col_idx, cell in enumerate(row):
                # Cell might be a list of paragraphs
                if isinstance(cell, list):
                    # Join all paragraphs in the cell
                    cell_text = " ".join(str(p).strip() for p in cell if p)
                else:
                    cell_text = str(cell).strip() if cell else ""

                cells.append(TableCell(row=row_idx, col=col_idx, content=cell_text))

        if not cells:
            return None

        # Calculate dimensions
        max_row = max(cell.row for cell in cells) if cells else 0
        max_col = max(cell.col for cell in cells) if cells else 0

        return TableBlock(
            id=self._generate_block_id(),
            type=BlockType.TABLE,
            content={},
            rows=max_row + 1,
            cols=max_col + 1,
            cells=cells,
            headers=headers,
            metadata=BlockMetadata(
                attributes={"section_index": section_idx, "table_index": table_idx}, confidence=0.95
            ),
        )

    def _process_header_footer(
        self, content: List, block_type: BlockType, section_id: str
    ) -> List[BaseBlock]:
        """Process headers or footers"""
        blocks = []

        if not content:
            return blocks

        # Flatten nested structure and extract text
        text_parts = []

        def extract_text(item):
            if isinstance(item, str):
                if item.strip():
                    text_parts.append(item.strip())
            elif isinstance(item, list):
                for subitem in item:
                    extract_text(subitem)

        extract_text(content)

        if text_parts:
            combined_text = " ".join(text_parts)
            blocks.append(
                BaseBlock(
                    id=self._generate_block_id(),
                    type=block_type,
                    content=combined_text,
                    metadata=BlockMetadata(attributes={"section": section_id}, confidence=0.95),
                )
            )

        return blocks

    def _process_footnotes(self, footnotes: List) -> List[BaseBlock]:
        """Process footnotes"""
        blocks = []

        for idx, footnote in enumerate(footnotes):
            if isinstance(footnote, str) and footnote.strip():
                blocks.append(
                    BaseBlock(
                        id=self._generate_block_id(),
                        type=BlockType.FOOTNOTE,
                        content=footnote.strip(),
                        metadata=BlockMetadata(
                            attributes={"footnote_index": idx + 1}, confidence=0.95
                        ),
                    )
                )
            elif isinstance(footnote, list):
                # Handle nested footnote structure
                text = " ".join(str(item) for item in footnote if item)
                if text.strip():
                    blocks.append(
                        BaseBlock(
                            id=self._generate_block_id(),
                            type=BlockType.FOOTNOTE,
                            content=text.strip(),
                            metadata=BlockMetadata(
                                attributes={"footnote_index": idx + 1}, confidence=0.95
                            ),
                        )
                    )

        return blocks

    def _process_endnotes(self, endnotes: List) -> List[BaseBlock]:
        """Process endnotes - similar to footnotes"""
        blocks = []

        for idx, endnote in enumerate(endnotes):
            if isinstance(endnote, str) and endnote.strip():
                blocks.append(
                    BaseBlock(
                        id=self._generate_block_id(),
                        type=BlockType.REFERENCE,
                        content=endnote.strip(),
                        metadata=BlockMetadata(
                            attributes={"endnote_index": idx + 1}, confidence=0.95
                        ),
                    )
                )

        return blocks

    def _extract_comments(self, docx_content) -> List[BaseBlock]:
        """Extract comments as blocks"""
        blocks = []

        if not docx_content.comments:
            return blocks

        for idx, comment in enumerate(docx_content.comments):
            if isinstance(comment, str) and comment.strip():
                blocks.append(
                    BaseBlock(
                        id=self._generate_block_id(),
                        type=BlockType.COMMENT,
                        content=comment.strip(),
                        metadata=BlockMetadata(
                            attributes={
                                "comment_index": idx + 1,
                                "comment_type": "document_comment",
                            },
                            confidence=0.95,
                        ),
                    )
                )
            elif isinstance(comment, dict):
                # Handle structured comment data
                text = comment.get("text", "")
                author = comment.get("author", "")
                if text:
                    blocks.append(
                        BaseBlock(
                            id=self._generate_block_id(),
                            type=BlockType.COMMENT,
                            content=text,
                            metadata=BlockMetadata(
                                attributes={
                                    "comment_index": idx + 1,
                                    "author": author,
                                    "comment_type": "document_comment",
                                },
                                confidence=0.95,
                            ),
                        )
                    )

        return blocks

    def _extract_images(self, docx_content) -> List[ImageBlock]:
        """Extract images as blocks"""
        blocks = []

        if not hasattr(docx_content, "images") or not docx_content.images:
            return blocks

        for filename, image_data in docx_content.images.items():
            # Create base64 data URI
            import base64

            image_base64 = base64.b64encode(image_data).decode("utf-8")

            # Detect image format from filename
            ext = Path(filename).suffix.lower()
            mime_type = {
                ".png": "image/png",
                ".jpg": "image/jpeg",
                ".jpeg": "image/jpeg",
                ".gif": "image/gif",
                ".bmp": "image/bmp",
                ".tiff": "image/tiff",
                ".svg": "image/svg+xml",
            }.get(ext, "image/png")

            data_uri = f"data:{mime_type};base64,{image_base64}"

            blocks.append(
                ImageBlock(
                    id=self._generate_block_id(),
                    type=BlockType.IMAGE,
                    content="",
                    src=data_uri,
                    alt=filename,
                    format=ext[1:] if ext else "unknown",
                    metadata=BlockMetadata(
                        attributes={"original_filename": filename, "embedded": True}, confidence=1.0
                    ),
                )
            )

        return blocks

    def _build_hierarchy(self, blocks: List[BaseBlock]) -> Optional[HierarchyNode]:
        """Build document hierarchy from heading blocks"""
        heading_blocks = [b for b in blocks if b.type == BlockType.HEADING]

        if not heading_blocks:
            # Try to build from other structural elements
            return self._build_default_hierarchy(blocks)

        # Create root node
        root = HierarchyNode(id="root", title="Document", level=0, block_id="root", children=[])

        # Build hierarchy
        stack = [root]

        for block in heading_blocks:
            # Add null check for metadata.attributes
            if block.metadata and block.metadata.attributes:
                level = block.metadata.attributes.get("level", 1)
            else:
                level = 1

            # Pop stack until we find parent level
            while len(stack) > level:
                stack.pop()

            # Create node
            node = HierarchyNode(
                id=f"h-{block.id}",
                title=block.content,
                level=level,
                block_id=block.id,
                parent_id=stack[-1].id if stack else None,
                breadcrumb=[n.title for n in stack] + [block.content],
            )

            # Add to parent
            if stack:
                stack[-1].children.append(node)

            # Update stack
            if len(stack) <= level:
                stack.append(node)
            else:
                stack[level] = node

        return root

    def _build_default_hierarchy(self, blocks: List[BaseBlock]) -> Optional[HierarchyNode]:
        """Build a default hierarchy when no headings are found"""
        root = HierarchyNode(id="root", title="Document", level=0, block_id="root", children=[])

        # Group by section if we have that info
        sections = {}
        for block in blocks:
            # Add null check for metadata.attributes
            if block.metadata and block.metadata.attributes:
                section_idx = block.metadata.attributes.get("section_index", 0)
            else:
                section_idx = 0
            if section_idx not in sections:
                sections[section_idx] = []
            sections[section_idx].append(block)

        # Create section nodes
        for section_idx in sorted(sections.keys()):
            if len(sections) > 1:  # Only create section nodes if multiple sections
                section_node = HierarchyNode(
                    id=f"section-{section_idx}",
                    title=f"Section {section_idx + 1}",
                    level=1,
                    block_id=f"section-{section_idx}",
                    parent_id="root",
                    breadcrumb=["Document", f"Section {section_idx + 1}"],
                )
                root.children.append(section_node)

        return root if root.children else None

    def _extract_full_text(self, blocks: List[BaseBlock]) -> str:
        """Extract all text content from blocks"""
        text_parts = []
        for block in blocks:
            if hasattr(block, "content") and isinstance(block.content, str):
                text_parts.append(block.content)
            elif block.type == BlockType.TABLE and hasattr(block, "cells"):
                # Extract table text
                for cell in block.cells:
                    if cell.content:
                        text_parts.append(cell.content)
        return "\n".join(text_parts)

    def _extract_keywords(self, docx_content) -> List[str]:
        """Extract keywords from document properties"""
        keywords = []

        if docx_content.core_properties:
            # Check for keywords in properties
            props = docx_content.core_properties
            if "keywords" in props and props["keywords"]:
                # Keywords might be comma-separated
                kw_string = props["keywords"]
                keywords = [k.strip() for k in kw_string.split(",") if k.strip()]

            # Also check subject
            if "subject" in props and props["subject"]:
                keywords.append(props["subject"])

        return keywords

    def _generate_block_id(self) -> str:
        """Generate unique block ID"""
        self.block_counter += 1
        return f"docx-block-{self.block_counter}"

    def _analyze_and_merge_tables(self, blocks: List[BaseBlock]) -> List[BaseBlock]:
        """Analyze tables for potential merging using Claude's logic"""
        # Import the Claude table merge analyzer if available
        try:
            from extractor.core.processors.claude_table_merge_analyzer import (
                ClaudeTableMergeAnalyzer,
            )

            analyzer = ClaudeTableMergeAnalyzer()
            return analyzer.analyze_and_merge_tables(blocks)
        except ImportError:
            logger.debug("Claude table merge analyzer not available, skipping table merge analysis")
            return blocks


if __name__ == "__main__":
    # Test the DOCX extractor

    # For testing, we'll create a simple test
    # In real usage, we'd have an actual DOCX file
    provider = DOCXProvider()

    # Since we can't easily create a DOCX file in memory for testing,
    # we'll just verify the class initializes correctly
    assert provider is not None
    assert provider.block_counter == 0

    print("✅ DOCX provider initialized successfully")
    print("Note: Full testing requires actual DOCX files")
