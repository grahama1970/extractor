"""Structured-document pipelines for non-PDF formats."""

from __future__ import annotations

import importlib.util
import json
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Type, Optional

from loguru import logger

from extractor.core.providers.docx import DOCXProvider
from extractor.core.providers.epub import EPUBProvider
from extractor.core.providers.html import HTMLProvider
from extractor.core.providers.pptx import PPTXProvider
from extractor.core.providers.rst import RSTProvider
from extractor.core.providers.spreadsheet import SpreadsheetProvider
from extractor.core.providers.xml import XMLProvider
from extractor.core.providers.markdown import MarkdownProvider
from extractor.core.schema.unified_document import (
    BaseBlock,
    BlockType,
    TableBlock,
    UnifiedDocument,
)
from extractor.pipeline.utils.unified_conversion import build_unified_document_from_reflow


@dataclass(frozen=True)
class StructuredPipelineMeta:
    format_name: str
    stage_prefix: str


STRUCTURED_PIPELINES: Dict[Type, StructuredPipelineMeta] = {
    HTMLProvider: StructuredPipelineMeta("html", "07_html_ingest"),
    DOCXProvider: StructuredPipelineMeta("docx", "07_docx_ingest"),
    PPTXProvider: StructuredPipelineMeta("pptx", "07_pptx_ingest"),
    SpreadsheetProvider: StructuredPipelineMeta("spreadsheet", "07_spreadsheet_ingest"),
    EPUBProvider: StructuredPipelineMeta("epub", "07_epub_ingest"),
    RSTProvider: StructuredPipelineMeta("rst", "07_rst_ingest"),
    XMLProvider: StructuredPipelineMeta("xml", "07_xml_ingest"),
    MarkdownProvider: StructuredPipelineMeta("markdown", "07_markdown_ingest"),
}

_FLATTEN_FN = None


def _load_flatten_function():
    global _FLATTEN_FN
    if _FLATTEN_FN is not None:
        return _FLATTEN_FN
    module_path = Path(__file__).resolve().parent / "steps" / "10_arangodb_exporter.py"
    spec = importlib.util.spec_from_file_location("pipeline_stage10", module_path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Unable to load Stage 10 module from {module_path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)  # type: ignore[attr-defined]
    _FLATTEN_FN = module.flatten_document_to_pdf_objects  # type: ignore[attr-defined]
    return _FLATTEN_FN


def _ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def _write_json(path: Path, payload: Dict) -> None:
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False))


def _text_from_block(block: BaseBlock) -> Optional[str]:
    content = block.content
    if isinstance(content, str):
        text = content.strip()
        if not text:
            return None
        if text.startswith("[[") and text.endswith("]]" ):
            return None
        if "[[" in text and "]]" in text:
            return None
        if text.startswith("----Image alt text----"):
            return None
        return text
    if isinstance(content, list):
        parts = [str(item).strip() for item in content if str(item).strip()]
        return " ".join(parts) or None
    if isinstance(content, dict):
        items = content.get("items")
        if isinstance(items, list) and items:
            return "\n".join(f"• {str(item).strip()}" for item in items if str(item).strip())
        for key in ("text", "content", "value"):
            value = content.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()
    return None


def _table_to_cells(table_block: TableBlock) -> Tuple[int, int, List[List[str]]]:
    cells = table_block.cells or []
    max_row = max((cell.row for cell in cells), default=-1)
    max_col = max((cell.col for cell in cells), default=-1)
    rows = max(table_block.rows or 0, max_row + 1)
    cols = max(table_block.cols or 0, max_col + 1)
    grid = [["" for _ in range(cols)] for _ in range(rows)]
    for cell in cells:
        try:
            grid[cell.row][cell.col] = cell.content or ""
        except IndexError:
            continue
    return rows, cols, grid


def _build_sections_from_blocks(unified_doc: UnifiedDocument) -> List[Dict]:
    sections: List[Dict] = []
    stack: List[Tuple[int, Dict]] = []
    current_section: Optional[Dict] = None
    default_title = unified_doc.metadata.title or "Document"

    def start_section(title: str, level: int) -> Dict:
        nonlocal current_section
        title = title or default_title
        while stack and stack[-1][0] >= level:
            stack.pop()
        parent_id = stack[-1][1]["id"] if stack else None
        section_id = f"section-{len(sections)}"
        section = {
            "id": section_id,
            "title": title,
            "level": level,
            "parent_id": parent_id,
            "reflow_status": "native",
            "paragraphs": [],
            "tables": [],
            "figures": [],
        }
        sections.append(section)
        stack.append((level, section))
        current_section = section
        return section

    for block in unified_doc.blocks:
        if block.type == BlockType.HEADING:
            level = 1
            if block.metadata and block.metadata.attributes:
                candidate = block.metadata.attributes.get("level")
                if isinstance(candidate, int) and candidate > 0:
                    level = candidate
                elif isinstance(candidate, str) and candidate.isdigit():
                    level = int(candidate)
            heading_text = _text_from_block(block) or default_title
            start_section(heading_text, level)
            continue

        if current_section is None:
            current_section = start_section(default_title, 1)

        if isinstance(block, TableBlock) or block.type == BlockType.TABLE:
            if isinstance(block, TableBlock):
                rows, cols, grid = _table_to_cells(block)
                table_title = None
                if isinstance(block.content, dict):
                    table_title = block.content.get("title")
                page_number = block.metadata.page_number if block.metadata else None
                sheet_attr = None
                try:
                    if block.metadata and block.metadata.attributes:
                        sheet_attr = block.metadata.attributes.get("sheet")
                except Exception:
                    sheet_attr = None
            else:
                content = block.content if isinstance(block.content, dict) else {}
                grid = content.get("cells") or []
                rows = content.get("rows") or len(grid)
                cols = content.get("cols") or (len(grid[0]) if grid else 0)
                table_title = content.get("title")
                page_number = content.get("page") or None
                if grid and rows and cols:
                    normalized_grid = []
                    for row in grid:
                        if isinstance(row, list):
                            normalized_grid.append([str(item) for item in row])
                        else:
                            normalized_grid.append([str(row)])
                    grid = normalized_grid
                else:
                    rows = cols = 0
            table_dict = {
                "table_id": block.id,
                "title": table_title,
                "rows": rows,
                "cols": cols,
                "cells": grid,
                "source": "structured",
                "page_number": page_number,
            }
            try:
                if 'sheet_attr' in locals() and sheet_attr:
                    table_dict["sheet"] = sheet_attr
            except Exception:
                pass
            # Preserve header rows if available on TableBlock
            try:
                headers = block.headers if isinstance(block, TableBlock) else None
                if headers:
                    table_dict["headers"] = headers
            except Exception:
                pass
            current_section["tables"].append(table_dict)
            continue

        if block.type in (BlockType.FIGURE, BlockType.IMAGE):
            data = block.content if isinstance(block.content, dict) else {}
            figure_dict = {
                "figure_id": block.id,
                "title": data.get("title"),
                "caption": data.get("caption") or data.get("description"),
                "image_path": data.get("image_path") or data.get("src"),
                "source": "structured",
                "page": block.metadata.page_number if block.metadata else None,
            }
            current_section["figures"].append(figure_dict)
            continue

        text = _text_from_block(block)
        if text:
            current_section.setdefault("paragraphs", []).append(text)

    if not sections:
        section = start_section(default_title, 1)
        paras: List[str] = []
        for block in unified_doc.blocks:
            txt = _text_from_block(block)
            if txt:
                paras.append(txt)
        section["paragraphs"].extend(paras)

    for section in sections:
        paras = section.pop("paragraphs", [])
        section["reflowed_text"] = "\n\n".join(paras)

    return sections


def _build_stage07_payload(
    unified_doc: UnifiedDocument,
    source_path: Path,
) -> Dict:
    sections = _build_sections_from_blocks(unified_doc)
    return {
        "timestamp": datetime.now().isoformat(),
        "source_files": {"sections": str(source_path)},
        "status": "Completed",
        "section_count": len(sections),
        "reflowed_sections": sections,
        "run_id": None,
        "errors_count": 0,
        "warnings_count": 0,
        "diagnostics": [],
        "timings": {},
        "resources": {},
        "unified_document": unified_doc.model_dump(by_alias=True, mode="json"),
    }


def run_structured_pipeline(
    provider_cls: Type,
    input_path: Path,
    results_dir: Path,
    *,
    stage_prefix: str,
    skip_export10: bool = True,
    skip_embeddings10: bool = True,
    fast_embeddings10: bool = True,
    auto_convert_mangled_docx: bool = True,
    fast_fallback_pdf: bool = True,
) -> Dict[str, Path]:
    provider = provider_cls()
    logger.info("Running {} pipeline for {}", provider_cls.__name__, input_path)

    mangled_info: Optional[Dict] = None
    if provider_cls is DOCXProvider:
        try:
            mangled, details = _docx_mangled_diagnose(input_path)
            mangled_info = {"mangled_docx": mangled, **details}
            if mangled and auto_convert_mangled_docx:
                logger.warning("DOCX appears mangled; attempting PDF fallback conversion.")
                pdf_path = _convert_docx_to_pdf(input_path, results_dir / input_path.stem / "converted")
                logger.info("Converted DOCX to PDF at {}", pdf_path)
                # Route to PDF pipeline and bypass structured extraction
                from extractor.pipeline import run_all as pdf_pipeline  # local import
                out_root = results_dir
                pdf_pipeline.run(
                    pdf=pdf_path,
                    results=out_root,
                    arango_db="pdf_knowledge_base_test",
                    session=None,
                    lean4_cli=None,
                    offline=True,
                    skip_llm03=True,
                    skip_descriptions06=True,
                    summary_only07=True,
                    skip_proving08=True,
                    skip_export10=True,
                    skip_embeddings10=True,
                    fast_embeddings10=True,
                    skip_graph11=True,
                    **({
                        'skip_tables05': True,
                        'skip_figures06': True,
                    } if fast_fallback_pdf else {}),
                    validate=False,
                    annotations_json=None,
                )
                pdf_stem = pdf_path.stem
                stage07_path = out_root / pdf_stem / "07_reflow_section" / "json_output" / "07_reflowed.json"
                stage10_path = out_root / pdf_stem / "10_arangodb_exporter" / "json_output" / "10_flattened_data.json"
                # Attach diagnostics file
                if mangled_info:
                    diag_path = out_root / pdf_stem / "07_reflow_section" / "json_output" / "07_mangled_docx_diag.json"
                    _write_json(diag_path, mangled_info)
                return {
                    "stage07": stage07_path,
                    "stage10_flattened": stage10_path,
                }
        except Exception as e:
            logger.warning("DOCX mangled detection failed: {}", e)

    unified_document = provider.extract_document(str(input_path))

    flatten = _load_flatten_function()

    run_dir = results_dir / input_path.stem
    stage07_dir = run_dir / stage_prefix / "json_output"
    stage10_dir = run_dir / "10_arangodb_exporter" / "json_output"
    _ensure_dir(stage07_dir)
    _ensure_dir(stage10_dir)

    stage07_payload = _build_stage07_payload(unified_document, input_path)
    if mangled_info is not None:
        stage07_payload.setdefault("diagnostics", [])
        stage07_payload["diagnostics"].append({"structured_pipeline": "docx_mangled_check", **mangled_info})
    # Merge provider format metadata (e.g., HTML generator) into document_metadata
    provider_meta = unified_document.metadata.format_metadata or {}
    doc_meta = {"source_files": {"sections": str(input_path)}}
    doc_meta.update(provider_meta)
    converted_unified = build_unified_document_from_reflow(
        sections=stage07_payload["reflowed_sections"],
        source_path=str(input_path),
        source_type=unified_document.source_type,
        document_metadata=doc_meta,
    )
    stage07_payload["unified_document"] = converted_unified.model_dump(
        by_alias=True, mode="json"
    )
    stage07_path = stage07_dir / "07_reflowed.json"
    _write_json(stage07_path, stage07_payload)
    logger.info("Wrote structured Stage 07 stub to {}", stage07_path)

    flattened = flatten(
        pipeline_data={
            "unified_document": converted_unified.model_dump(by_alias=True, mode="json"),
            "source_files": {"sections": str(input_path)},
        },
        summaries_data={"summaries": []},
        skip_embeddings=skip_embeddings10,
        fast_embeddings=fast_embeddings10,
    )

    stage10_path = stage10_dir / "10_flattened_data.json"
    stage10_path.write_text(json.dumps(flattened, indent=2))
    logger.info("Flattened document to {} (objects={})", stage10_path, len(flattened))

    # Stage 11 (offline JSON edges) — always emit graph edges JSON for structured runs
    try:
        stage11_dir = run_dir / "11_arango_create_graph" / "json_output"
        _ensure_dir(stage11_dir)
        bundle = stage11_dir / "_bundle.json"
        bundle.write_text(json.dumps({"documents": flattened}, indent=2))
        # Import and call debug_bundle to produce 11_graph_edges.json
        mod_path = Path(__file__).resolve().parent / "steps" / "11_arango_create_graph.py"
        spec = importlib.util.spec_from_file_location("pipeline_stage11", mod_path)
        if spec and spec.loader:
            module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(module)  # type: ignore
            # stage11 debug_bundle signature: (bundle: Path, output_dir: Path, k_neighbors: int, similarity_threshold: float)
            try:
                module.debug_bundle(bundle=bundle, output_dir=run_dir, k_neighbors=10, similarity_threshold=0.55)  # type: ignore[attr-defined]
            except TypeError:
                # Backward‑compat for different signature
                module.debug_bundle(bundle=bundle, output_dir=run_dir)  # type: ignore[attr-defined]
        else:
            logger.warning("Stage 11 module not loadable; skipping graph edges for structured run")
    except Exception as e:
        logger.warning("Stage 11 (offline edges) failed for structured run: {}", e)

    if not skip_export10:
        logger.warning("Structured pipelines do not yet export directly to ArangoDB; skipping.")

    return {
        "stage07": stage07_path,
        "stage10_flattened": stage10_path,
        "stage11_edges": (run_dir / "11_arango_create_graph" / "json_output" / "11_graph_edges.json"),
    }


def _docx_mangled_diagnose(input_path: Path) -> Tuple[bool, Dict[str, object]]:
    """Heuristically determine if a DOCX likely came from a PDF export (layout-only).

    Signals:
      - Very low fraction of Heading-styled paragraphs
      - No numbering (w:numPr) usage
      - No real tables in python-docx
      - High image-to-paragraph ratio
    """
    try:
        from docx import Document as PythonDocxDocument  # type: ignore
    except Exception:
        return False, {"error": "python-docx-unavailable"}

    doc = PythonDocxDocument(str(input_path))
    total_pars = len(doc.paragraphs) or 1
    heading_pars = sum(1 for p in doc.paragraphs if (getattr(p.style, "name", "") or "").lower().startswith("heading"))

    # Numbering detection via direct element
    numbered = 0
    for p in doc.paragraphs:
        try:
            pPr = p._p.pPr  # type: ignore[attr-defined]
            if pPr is not None and pPr.numPr is not None:  # type: ignore[attr-defined]
                numbered += 1
        except Exception:
            continue

    tables_count = len(doc.tables)
    # inline_shapes is not exposed; approximate via pictures in runs
    image_like = 0
    for p in doc.paragraphs:
        for r in p.runs:
            try:
                if r._r.drawing is not None:  # type: ignore[attr-defined]
                    image_like += 1
            except Exception:
                continue

    heading_ratio = heading_pars / max(total_pars, 1)
    numbered_ratio = numbered / max(total_pars, 1)
    image_ratio = image_like / max(total_pars, 1)

    # Simple thresholding
    # Strategy: if there are no real tables at all, treat as mangled for our use case.
    mangled = tables_count == 0 or (
        heading_ratio < 0.01 and numbered_ratio == 0 and image_ratio > 0.02
    )
    details = {
        "total_paragraphs": total_pars,
        "heading_paragraphs": heading_pars,
        "heading_ratio": round(heading_ratio, 4),
        "numbered_paragraphs": numbered,
        "numbered_ratio": round(numbered_ratio, 4),
        "tables_count": tables_count,
        "image_like_runs": image_like,
        "image_ratio": round(image_ratio, 4),
    }
    return mangled, details


def _convert_docx_to_pdf(input_path: Path, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = out_dir / (input_path.stem + ".pdf")
    # Try soffice
    import shutil, subprocess
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            cmd = [soffice, "--headless", "--convert-to", "pdf:writer_pdf_Export", "--outdir", str(out_dir), str(input_path)]
            subprocess.run(cmd, check=True)
            if pdf_path.exists():
                return pdf_path
        except Exception:
            pass
    # Fallback pandoc
    pandoc = shutil.which("pandoc")
    if pandoc:
        try:
            cmd = [pandoc, str(input_path), "-o", str(pdf_path)]
            subprocess.run(cmd, check=True)
            if pdf_path.exists():
                return pdf_path
        except Exception:
            pass
    raise RuntimeError("DOCX→PDF conversion failed (soffice/pandoc not available or failed)")


def run_html_pipeline(
    input_path: Path,
    results_dir: Path,
    *,
    skip_export10: bool = True,
    skip_embeddings10: bool = True,
    fast_embeddings10: bool = True,
) -> Dict[str, Path]:
    meta = STRUCTURED_PIPELINES[HTMLProvider]
    return run_structured_pipeline(
        HTMLProvider,
        input_path,
        results_dir,
        stage_prefix=meta.stage_prefix,
        skip_export10=skip_export10,
        skip_embeddings10=skip_embeddings10,
        fast_embeddings10=fast_embeddings10,
    )


FORMAT_TO_PROVIDER: Dict[str, Type] = {
    meta.format_name: provider_cls for provider_cls, meta in STRUCTURED_PIPELINES.items()
}
