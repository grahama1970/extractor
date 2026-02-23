from docx import Document
from pathlib import Path


def generate_docx():
    doc = Document()
    doc.add_heading("Structured DOCX Test", 0)

    doc.add_heading("Section 1: Introduction", level=1)
    doc.add_paragraph(
        "This is a test document to verify the multi-modal capabilities of the extractor. It contains structured headers, lists, and tables."
    )

    doc.add_heading("Section 2: Requirements", level=1)
    doc.add_paragraph(
        "REQ-001: The system shall extract text from DOCX files.", style="List Bullet"
    )
    doc.add_paragraph("REQ-002: The system shall preserve structure.", style="List Bullet")

    doc.add_heading("Section 3: Tabular Data", level=1)
    table = doc.add_table(rows=3, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "ID"
    hdr_cells[1].text = "Value"

    row1 = table.rows[1].cells
    row1[0].text = "1"
    row1[1].text = "Alpha"

    row2 = table.rows[2].cells
    row2[0].text = "2"
    row2[1].text = "Beta"

    output_dir = Path("tools/tasks_loop/fixtures/structured_docx")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "source.docx"

    doc.save(output_path)
    print(f"Generated DOCX: {output_path}")


if __name__ == "__main__":
    generate_docx()
